from __future__ import annotations

import math
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Conflict_Aware_Timetable_Management_System_Documentation_Sample_Format.docx"


PROJECT_OVERVIEW = {
    "title": "CONFLICT AWARE TIMETABLE MANAGEMENT SYSTEM",
    "subtitle": "Comprehensive Project Documentation",
    "department": "Department of Computer Science and Engineering",
    "institution": "Vikrama Simhapuri University Inspired Academic Workflow",
    "tech_stack": "Python, Django, MySQL, HTML, CSS, Bootstrap and JavaScript",
}


CHAPTERS = [
    {
        "number": 1,
        "title": "INTRODUCTION",
        "start_page": 9,
        "intro": "This chapter introduces the academic problem that motivated the development of the Conflict Aware Timetable Management System and explains the institutional setting in which the application operates.",
        "sections": [
            {
                "number": "1.1",
                "title": "Problem Statement",
                "paragraphs": [
                    "Higher education institutions manage several interconnected activities at the same time: academic planning, faculty assignment, student access, syllabus circulation, question paper archiving and timetable publication. When these tasks are handled manually or through isolated spreadsheets, the probability of conflict rises sharply. The same faculty member may be assigned to overlapping slots, laboratory sessions may consume fragmented hours, and departments may struggle to maintain a single current timetable that everyone can trust.",
                    "The present project addresses this problem by centralizing timetable and academic support functions inside a single web application. The implemented system does not treat timetable generation as a standalone spreadsheet exercise. Instead, it connects user roles, subject definitions, faculty ownership, semester data and resource management so that the published schedule reflects the operational state of the department rather than a disconnected draft prepared in isolation.",
                    "A major institutional concern captured in this project is role separation. The main assistant monitors the system at a university level, department assistants manage departmental master data, faculty members observe their assigned workload, and students consume the final outputs. This separation reduces accidental modifications, distributes responsibility clearly and makes the application suitable for real academic administration where accountability is as important as convenience.",
                ],
                "bullets": [
                    "Manual timetable preparation is slow and conflict-prone.",
                    "Distributed files make version control difficult across departments.",
                    "Departments need a secure role-based workflow for data entry and review.",
                    "Students and faculty require a single trusted source for timetable and resources.",
                ],
            },
            {
                "number": "1.2",
                "title": "Objectives of the Project",
                "paragraphs": [
                    "The primary objective of the project is to design and implement a role-based academic management platform capable of generating department-wise timetables while reducing scheduling conflicts. The system is expected to support the full flow from master-data creation to timetable viewing, which means that faculty records, subject definitions and semester selection must exist inside the same controlled environment as the timetable generator.",
                    "A second objective is to establish a disciplined administrative workflow. The project creates a main assistant account during first-run setup, allows that role to pre-register department assistants, and then delegates operational data management to the appropriate department. This model is reflected consistently across the application and prevents a central monitoring user from bypassing departmental ownership rules during normal data-entry operations.",
                    "The third objective is to support academic resource continuity beyond timetable generation. Because students often need syllabi and old question papers alongside the weekly schedule, the application integrates a resource-management module for uploads, browsing and filtering. This transforms the project from a narrow timetable utility into a more meaningful academic service platform.",
                ],
                "table": {
                    "headers": ["Objective", "Meaning in the Implemented System"],
                    "rows": [
                        ["Conflict reduction", "Avoid duplicate lab blocks, repeated same-day theory placement and uncontrolled faculty load."],
                        ["Role-based control", "Separate the monitoring role from departmental data-entry responsibilities."],
                        ["Unified access", "Keep timetable, users, syllabus and old papers in one authenticated portal."],
                        ["Database reliability", "Persist institutional data in MySQL for multi-user academic deployment."],
                    ],
                },
            },
            {
                "number": "1.3",
                "title": "Scope of the Project",
                "paragraphs": [
                    "The scope of the current implementation covers university-style departmental timetable administration for a weekly six-day teaching pattern with a dedicated lunch interval and six teaching slots per day. The application is especially suitable for departments that operate multiple semesters in parallel and need a quick way to visualize schedules semester-wise while still tracking faculty workload across the complete academic year.",
                    "Within this scope, the system supports custom users, first-run bootstrap, self-registration through pre-authorized email addresses, faculty management, subject management, timetable generation, workload reporting, syllabus uploads, old question paper uploads and basic password reset workflows. It also includes a Django admin backend and a seed-data command, which improves maintainability and helps demonstrate how the system behaves during initial institutional setup.",
                    "The current version does not attempt to solve every enterprise scheduling case. It does not integrate classroom capacity, transportation constraints, attendance analytics or automated email delivery. However, the architecture and module separation create a strong base for such enhancements. As a result, the documented system should be viewed as a practical academic-management core that can be extended in future releases.",
                ],
            },
            {
                "number": "1.4",
                "title": "Existing System and Proposed System",
                "paragraphs": [
                    "In many institutions, the existing process depends on spreadsheets, message groups and manual approvals. These tools provide short-term flexibility but create long-term coordination problems because data is repeated in multiple files and there is no guaranteed synchronization between timetable edits and faculty or subject changes. Security is also weak because anyone with a file copy may create an unofficial revision that circulates among students.",
                    "The proposed system replaces that fragmented model with a web application that recognizes institutional roles and stores academic records in a relational database. Instead of editing a timetable directly as a static table, department assistants define faculty, subjects and academic years first, and the system uses those inputs to generate structured output. The result is a process that is easier to audit, easier to reproduce and more resilient when multiple users participate in academic administration.",
                    "The move from SQLite-oriented development to MySQL-backed deployment further strengthens the proposed system. MySQL supports multi-user access more naturally, aligns better with production hosting environments and improves the readiness of the project for academic deployment where several administrative users may work concurrently. This makes the documented implementation more aligned with real institutional expectations than a purely local prototype.",
                ],
                "table": {
                    "headers": ["Dimension", "Existing Manual Approach", "Proposed Web System"],
                    "rows": [
                        ["Data storage", "Spreadsheets and scattered files", "Centralized MySQL database"],
                        ["Role control", "Informal or file-based", "Authenticated role-based access"],
                        ["Timetable creation", "Manual drafting", "Rule-aware automated generation"],
                        ["Resource sharing", "Separate document circulation", "Integrated upload and access module"],
                    ],
                },
            },
        ],
    },
    {
        "number": 2,
        "title": "LITERATURE REVIEW",
        "start_page": 13,
        "intro": "This chapter summarizes conceptual and technical ideas relevant to timetable automation, academic resource management and role-based information systems.",
        "sections": [
            {
                "number": "2.1",
                "title": "Academic Timetable Automation in Practice",
                "paragraphs": [
                    "Timetable automation has long been recognized as a constrained scheduling problem in which limited slots, shared faculty and semester-specific teaching loads must be arranged without producing avoidable clashes. Many academic systems begin by modelling days, periods, rooms, subjects and instructors as scheduling entities. Even when the final algorithm is simple, the quality of the data model determines whether the generated result is meaningful or merely random.",
                    "The implemented project follows a practical interpretation of this literature by focusing on the constraints most relevant to departmental weekly planning. Subjects are grouped by semester and academic year, laboratories are treated as longer blocks, and scheduling decisions are spread across the week to avoid repetitive placement. This does not attempt to replicate complex heuristic research in full, but it does apply the core principle that conflict reduction should be driven by structured rules rather than by ad hoc copying and pasting.",
                    "A second lesson from existing research is that timetable tools fail when they ignore the surrounding administrative ecosystem. Schedules are not independent from user authorization, academic-year tracking or faculty ownership. The present project reflects that lesson by embedding timetable generation within a broader system where users, faculty codes, subject assignment and resource circulation are all first-class concerns.",
                ],
            },
            {
                "number": "2.2",
                "title": "Role-Based Information Systems in Education",
                "paragraphs": [
                    "Educational software is most effective when access rights reflect the real distribution of academic responsibility. Administrators typically require broader visibility than operational staff, yet unrestricted edit permission often produces governance problems. Role-based access control addresses this issue by linking permissions to institutional responsibility rather than to technical convenience alone.",
                    "The project under study implements this concept in a particularly clear form. The main assistant is intentionally limited to a monitor and coordinator role, while department assistants own local operational data. Faculty and students consume timetable and resource outputs without being allowed to modify master records. This design choice is a functional feature rather than a limitation because it reduces the risk of central overreach and mirrors how academic authority is often distributed in practice.",
                    "The associated benefit is traceability. Because each role has a distinct interface and a defined operational boundary, it becomes easier to explain who is responsible for a missing faculty record, an incorrect subject mapping or an unuploaded syllabus. In a university environment, this accountability is a major requirement for sustainable software adoption.",
                ],
                "bullets": [
                    "Main assistant: monitoring and department-assistant onboarding.",
                    "Department assistant: faculty, subjects, users and resources for one department.",
                    "Faculty: timetable viewing with personal highlight support.",
                    "Student: timetable and academic-resource access.",
                ],
            },
            {
                "number": "2.3",
                "title": "Web-Based Academic Resource Management",
                "paragraphs": [
                    "Modern academic systems increasingly integrate supporting documents with operational data. Students expect not only a timetable but also immediate access to semester-wise syllabi and previous examination material. When these documents are distributed outside the timetable platform, students must depend on multiple channels and may encounter outdated or incomplete files.",
                    "The implemented system addresses this need through dedicated models for syllabus and old question papers. Both resources are stored with department, academic year and semester metadata so that retrieval remains meaningful within a specific academic context. This is more useful than storing files in a flat repository because the metadata creates a navigable academic memory that students and faculty can filter according to program needs.",
                    "From a documentation perspective, the resource module is important because it proves that the application is not a single-purpose prototype. It supports a broader academic workflow, improves student convenience and makes the final project documentation more representative of a deployable department portal rather than of a purely algorithmic demonstration.",
                ],
            },
        ],
    },
    {
        "number": 3,
        "title": "METHODOLOGY",
        "start_page": 16,
        "intro": "This chapter explains how the project was planned, structured, implemented and refined from requirements to executable modules.",
        "sections": [
            {
                "number": "3.1",
                "title": "Requirement Identification",
                "paragraphs": [
                    "Requirement identification for the project began by separating institutional concerns into three categories: operational master data, generated outputs and user access boundaries. Master data includes faculty records, subjects, academic years, semester values and pre-registered users. Generated outputs include the timetable matrix and faculty workload summary. Access boundaries govern who is allowed to create, review or consume each type of information.",
                    "This decomposition was valuable because it prevented the system from being designed around the timetable alone. By identifying the surrounding data dependencies early, the project ensured that the final generator would have reliable inputs and that schedule visibility would remain tied to authenticated users. It also highlighted the need for a first-run setup process so that the very first administrative account could be created before normal pre-registration rules became active.",
                    "The resulting requirement list was then translated into application modules. Accounts handles identity and onboarding. Core handles faculty, subjects, timetable generation and resources. The project configuration layer defines routes, installed applications, media handling and database settings. This modular decomposition provides both conceptual clarity and maintainable implementation boundaries.",
                ],
            },
            {
                "number": "3.2",
                "title": "Architectural Planning",
                "paragraphs": [
                    "Architectural planning followed the Django application pattern because the project needed built-in authentication support, ORM-backed persistence, HTML rendering and administrative tooling. The use of separate apps for accounts and core makes the design easier to understand and aligns the code structure with functional ownership. Each app encapsulates its models, forms, views and supporting logic, which reduces accidental coupling and improves future extensibility.",
                    "The selected architecture also supports a layered reasoning style. Models define academic entities and relationships. Forms validate user input and enforce data quality in the interface layer. Views orchestrate request handling, authorization and persistence. Templates present role-specific screens that adapt to the logged-in user. Because these layers remain visible in the codebase, the architecture is suitable for academic explanation and not only for execution.",
                    "The migration from SQLite to MySQL required no redesign of the architecture itself, which demonstrates a sound abstraction boundary. Django’s ORM hides most vendor-specific details while still allowing the settings layer to declare MySQL-specific configuration such as engine choice, host, port and initialization behavior. This is a useful methodological result because it shows that portability was preserved at the application layer.",
                ],
                "table": {
                    "headers": ["Layer", "Main Responsibility", "Representative Files"],
                    "rows": [
                        ["Configuration", "Global settings and routing", "vsu_timetable/settings.py, urls.py"],
                        ["Accounts app", "Identity, registration and password reset", "accounts/models.py, forms.py, views.py"],
                        ["Core app", "Faculty, subjects, resources and timetable logic", "core/models.py, views.py, timetable.py"],
                        ["Templates", "Role-specific user interface", "templates/accounts/*, templates/core/*"],
                    ],
                },
            },
            {
                "number": "3.3",
                "title": "Implementation Workflow",
                "paragraphs": [
                    "Implementation proceeded incrementally. The custom user model and registration flow were established first because almost every other module depends on reliable user identity and role metadata. From there, faculty and subject models were created so that academic records could be entered in department context. Timetable generation was added only after those records existed, which kept the algorithm grounded in realistic data rather than in placeholder inputs.",
                    "The workflow then expanded into user experience details such as dashboard statistics, department-scoped management views, resource uploads and print-friendly timetable output. In parallel, administrative safeguards were added through decorators and mixins so that each view respected the system’s intended authority structure. This sequencing reflects a practical software methodology in which functional completeness and governance evolve together instead of being treated as separate concerns.",
                    "Finally, the data layer was adapted for MySQL deployment. Requirements were updated to include PyMySQL and cryptography support, the settings module was reconfigured to read MySQL connection values from the environment, and migration behavior was validated against the relational schema. This completed the transition from local prototype assumptions to a more deployment-ready academic application.",
                ],
            },
            {
                "number": "3.4",
                "title": "Validation and Refinement Approach",
                "paragraphs": [
                    "Validation in this project is both technical and institutional. Technical validation asks whether the forms, models, migrations and timetable algorithm behave correctly. Institutional validation asks whether the roles, workflows and restrictions reflect the intended administrative process. A timetable generator that produces cells quickly but ignores departmental control boundaries would not be successful in a real academic environment.",
                    "The project therefore uses multiple refinement signals. Model-level validation prevents duplicate assistant roles in the same logical slot. View-level checks deny unauthorized operations to users outside the permitted role. Template-level cues communicate access limits and workflow expectations clearly to the user. A small automated test suite exercises the account-switching flow, and manual runs of migrations and server startup confirm that the system behaves coherently on the chosen MySQL backend.",
                    "This combination of refinement methods is appropriate for a university project. It demonstrates coding ability, database awareness, user-interface reasoning and operational thinking at the same time. The methodology chapter therefore captures not just how the software was coded, but how the system was shaped to fit academic administration as an end-to-end workflow.",
                ],
            },
        ],
    },
    {
        "number": 4,
        "title": "SYSTEM REQUIREMENTS AND SPECIFICATION",
        "start_page": 20,
        "intro": "This chapter lists the software, hardware and environmental assumptions needed for development, execution and basic academic deployment.",
        "sections": [
            {
                "number": "4.1",
                "title": "Software Requirements",
                "paragraphs": [
                    "The software stack of the project is intentionally modern but compact. The backend uses Python and Django 4.2, which provide the authentication framework, templating engine, migration system and ORM abstraction required for a structured academic management platform. Frontend pages are rendered through Django templates enhanced with HTML, CSS, Bootstrap classes and lightweight JavaScript for dynamic forms and interactive state hints.",
                    "Database persistence now targets MySQL through Django’s MySQL backend and PyMySQL compatibility. This change is significant because it moves the system closer to a deployment-ready environment and allows multiple institutional users to rely on a central relational store rather than on a single-file prototype database. The cryptography package complements this setup by supporting MySQL authentication modes used on modern servers.",
                    "Additional software assumptions include a web browser for end-user access, a Python package environment for dependency management and a filesystem area for media uploads. Because syllabus and question papers are stored as uploaded files, media configuration and directory permissions are part of the practical runtime specification rather than optional extras.",
                ],
                "table": {
                    "headers": ["Component", "Specification", "Purpose"],
                    "rows": [
                        ["Programming language", "Python 3.x", "Application logic and command execution"],
                        ["Web framework", "Django 4.2.x", "Authentication, ORM, routing and templates"],
                        ["Database", "MySQL", "Persistent multi-user academic data storage"],
                        ["Driver packages", "PyMySQL and cryptography", "MySQL connectivity and authentication support"],
                        ["Frontend toolkit", "HTML, CSS, Bootstrap, JavaScript", "Responsive user interface"],
                    ],
                },
            },
            {
                "number": "4.2",
                "title": "Hardware Requirements",
                "paragraphs": [
                    "The project does not require specialized hardware, which makes it accessible for university departments with modest infrastructure. A standard desktop or laptop capable of running Python, MySQL and a development or lightweight hosting server is sufficient for experimentation and controlled deployment. This low hardware barrier is valuable in academic settings where software practicality often determines adoption more than algorithmic sophistication alone.",
                    "For a single-department or demonstration installation, a mid-range machine with adequate memory, stable storage and basic network connectivity is enough. If the platform grows to support multiple departments with frequent file uploads and concurrent access, the same architecture can move to a stronger server without changing the application design. The database engine and media directories would simply be hosted on a more capable environment.",
                    "The most important hardware consideration is reliability rather than raw performance. Because the system stores user identities, timetable data and academic files, backup strategy and disk integrity matter more than processor-intensive computation. The scheduling algorithm is lightweight compared with the value of the institutional information it manages.",
                ],
            },
            {
                "number": "4.3",
                "title": "Functional and Non-Functional Specification",
                "paragraphs": [
                    "Functionally, the system must allow first-run setup, login, logout, self-registration for pre-authorized users, password reset, faculty management, subject management, user pre-registration, timetable generation, workload visualization and academic-resource upload and retrieval. Each of these operations is present in the implemented codebase and is exposed through dedicated templates or administration screens.",
                    "Non-functional expectations are equally important. The application must remain understandable to administrative users, maintain data segregation between departments, provide reasonably quick response for normal academic workloads and support recoverable database-backed persistence. The move to MySQL strengthens the persistence requirement, while Django’s modular app structure supports maintainability and future enhancement.",
                    "Security and governance appear as implicit non-functional requirements. The system must prevent unauthorized editing, protect the uniqueness of assistant roles and avoid department cross-access for ordinary users. These qualities are enforced through model validation, filtered querysets, role-aware views and mixins that explicitly deny access when a user attempts an operation outside the intended authority boundary.",
                ],
            },
        ],
    },
    {
        "number": 5,
        "title": "SYSTEM ANALYSIS AND ARCHITECTURE",
        "start_page": 23,
        "intro": "This chapter explains how the implemented modules cooperate and how information flows from authentication to timetable publication.",
        "sections": [
            {
                "number": "5.1",
                "title": "High-Level Architecture",
                "paragraphs": [
                    "At a high level, the system can be understood as three coordinated layers. The configuration layer defines the Django project environment, global routes, static and media handling, and database behavior. The accounts layer manages user identity, registration and password recovery. The core layer manages domain entities such as faculty, subjects, syllabi, question papers and the timetable engine. Templates then present the output of these layers to role-specific users.",
                    "This high-level design supports clarity and containment. Changes to user rules can be made largely inside the accounts app without forcing timetable code to change. Similarly, changes to timetable logic can occur inside the core app while leaving the authentication architecture intact. Such separation is valuable in academic software because institutional rules and scheduling rules often evolve at different speeds.",
                    "The architecture also demonstrates a balanced use of Django conventions and project-specific logic. Standard framework mechanisms are used for routing, forms and authentication, while custom code is reserved for academic concerns such as assistant uniqueness, department scoping and conflict-aware timetable generation. This blend increases maintainability because the framework handles generic plumbing while the custom modules focus on domain meaning.",
                ],
            },
            {
                "number": "5.2",
                "title": "Role-Based Control Flow",
                "paragraphs": [
                    "Role-based control begins at login and continues throughout the navigation flow. If no users exist, the platform redirects to a first-run setup screen so that the inaugural main assistant can be created safely. Once users are present, the login page becomes the shared entry point for sign-in and self-registration. Registration itself is constrained by pre-registration rules, which means the platform only admits users whose email addresses have been authorized by an appropriate assistant.",
                    "After authentication, the dashboard changes according to the user’s role and department. The main assistant sees a monitor-oriented summary and can pre-register department assistants. A department assistant sees management cards for subjects, faculty, users and resources within one department only. Faculty and students see consumption-oriented flows that focus on timetable and resource visibility. This dynamic control flow is essential because it ensures that the same codebase behaves like several carefully limited academic tools instead of one all-powerful interface.",
                    "The enforcement mechanism is not merely visual. Decorators and mixins verify role conditions before the target view executes. Department filtering is then applied inside the view logic itself, so even if a user manipulates a query string manually, the application still restricts the resulting queryset. This layered protection is a strong architectural property for academic software.",
                ],
                "table": {
                    "headers": ["Role", "Primary Access", "Restricted From"],
                    "rows": [
                        ["Main Assistant", "Monitoring, department assistant onboarding, Django admin visibility", "Department-level data entry screens"],
                        ["Department Assistant", "Subjects, faculty, users, resources and timetable in one department", "Other departments and self-deletion loopholes"],
                        ["Faculty", "Timetable viewing and resource access", "Administrative editing screens"],
                        ["Student", "Timetable and academic-resource access", "Administrative editing screens"],
                    ],
                },
            },
            {
                "number": "5.3",
                "title": "Module Interaction",
                "paragraphs": [
                    "Module interaction in the project is intentionally straightforward. The accounts module owns user-related information and exposes current user role and department data that the core module consumes during filtering and authorization. The core module then uses that information to decide which records to show, which operations to permit and which timetable or resource subset is relevant to the current session.",
                    "A good example is the timetable page. The request enters the core timetable view, which first determines whether the user is a super-level assistant or a department-scoped user. It then derives the applicable department filter, collects subject data for the selected academic year and calls the generation function. Finally, the rendered template highlights personal faculty slots when the logged-in user is a faculty member. This flow uses identity information, domain data and presentation rules in a tightly coordinated but still readable manner.",
                    "Another example is self-registration. The registration form validates email against pre-registered records, applies assistant-uniqueness rules when necessary, creates the final user and marks the pre-registration slot as consumed. Here again, model data, form validation and view orchestration work together to produce a safe institutional workflow.",
                ],
            },
            {
                "number": "5.4",
                "title": "Navigation and User Experience Design",
                "paragraphs": [
                    "The user interface adopts a simple dashboard-first design. Important actions are presented as cards with strong role cues, and context banners explain what each role can or cannot do. This is especially valuable in a university project because many users are administrators or faculty members rather than technical specialists. The interface therefore aims to reduce ambiguity before the user even attempts an operation.",
                    "The timetable page is optimized for both viewing and printing. It presents a structured grid with morning and afternoon slots, a lunch separator and faculty workload data beneath the main matrix. The HTML and CSS deliberately use an Excel-like presentation style for the printable table so that exported output remains familiar to academic staff who traditionally work with spreadsheets and printed schedules.",
                    "User-experience design is also visible in small workflow details. The manage users page explains how pre-registration works. The manage resources page supports dynamic academic-year selection. Dashboard banners reflect whether the user is a monitoring assistant, department assistant, faculty member or student. These interface decisions are not cosmetic only; they actively support institutional correctness by making role boundaries visible and intuitive.",
                ],
            },
        ],
    },
    {
        "number": 6,
        "title": "DATABASE DESIGN AND MYSQL IMPLEMENTATION",
        "start_page": 27,
        "intro": "This chapter documents the persistent data model and the transition from SQLite-oriented development to MySQL-backed execution.",
        "sections": [
            {
                "number": "6.1",
                "title": "Entity Overview",
                "paragraphs": [
                    "The persistent design of the system revolves around a small but meaningful set of entities. User represents registered people with role, department and academic attributes. PreRegisteredUser stores authorization slots for future self-registration. PasswordResetToken captures recovery tokens and their validity state. Faculty, Subject, Syllabus and OldQuestionPaper model the academic objects required for timetable generation and resource distribution.",
                    "This entity set is sufficient because each model carries explicit academic meaning. Faculty exists independently so that subjects can be assigned to a named instructor with a stable code. Subject includes year, semester, hours per week and departmental context, which gives the timetable algorithm everything it needs to construct a weekly schedule. Resource entities preserve uploaded files together with metadata that supports student retrieval by department, year and semester.",
                    "An important feature of the design is that many institutional rules are represented directly in the schema. Unique-together constraints, protected foreign keys and default values reduce the chance of meaningless records entering the system. As a result, the database is not merely a storage container; it is an active participant in maintaining academic consistency.",
                ],
            },
            {
                "number": "6.2",
                "title": "User and Access-Control Tables",
                "paragraphs": [
                    "The custom user model is central to the database design because it extends the default authentication concept with academic roles and department ownership. The fields email, full_name, department, role, semester and academic_year make it possible to use one identity system for assistants, faculty and students without losing the distinctions needed for authorization and reporting.",
                    "PreRegisteredUser complements this design by representing a pending admission slot rather than a final account. This table is essential for controlled self-registration because it allows assistants to authorize users before those users choose passwords. PasswordResetToken then supports account recovery with limited validity periods and one-time consumption behavior. Together, these tables create a structured identity subsystem that is more appropriate for academic administration than a simple open-registration model.",
                    "Database rules reinforce the access model. The code validates assistant uniqueness by department, distinguishes a main assistant through a blank department convention, and preserves registration state for pre-registered records. These behaviors ensure that institutional identity logic survives across sessions and is not left to interface conventions alone.",
                ],
                "table": {
                    "headers": ["Entity", "Primary Purpose", "Important Fields"],
                    "rows": [
                        ["User", "Registered platform account", "email, full_name, role, department, semester, academic_year"],
                        ["PreRegisteredUser", "Authorized future registration slot", "email, role, department, academic_year, registered"],
                        ["PasswordResetToken", "Temporary recovery token", "user, token, created_at, used"],
                    ],
                },
            },
            {
                "number": "6.3",
                "title": "Academic Domain Tables",
                "paragraphs": [
                    "The academic domain layer is intentionally normalized around faculty and subjects. Faculty records hold a department, a short code and a display name. Subjects then reference a faculty record through a protected foreign key, ensuring that instructional ownership is explicit. The subject model also stores year, semester, code, name and hours per week, which together define the teaching load that must be scheduled.",
                    "Resource management extends the domain through Syllabus and OldQuestionPaper tables. These models are not directly involved in timetable generation, but they are crucial for the portal’s academic usefulness. Both include department, academic year and semester metadata so that students see material in a meaningful academic grouping rather than as an unstructured file repository.",
                    "The design choice to store uploaded_by as text rather than as a strict foreign key keeps resource records robust even if user accounts change later. Although a stricter relational link might be possible in future versions, the present design balances traceability and operational simplicity well for a university department portal.",
                ],
                "table": {
                    "headers": ["Entity", "Relationship Notes", "Constraint Highlights"],
                    "rows": [
                        ["Faculty", "Owned by one department", "Department and code are unique together"],
                        ["Subject", "Belongs to one faculty record", "Department, year, semester and code are unique together"],
                        ["Syllabus", "File resource by semester", "One department-year-semester combination per syllabus"],
                        ["OldQuestionPaper", "File resource by subject and exam context", "Ordered for retrieval by department, year and semester"],
                    ],
                },
            },
            {
                "number": "6.4",
                "title": "MySQL Configuration and Migration Strategy",
                "paragraphs": [
                    "The project was adapted to use MySQL as the primary database engine in place of SQLite. This required changes in the Django settings file to declare the MySQL backend and to read database name, user, password, host and port from environment variables. Sensible defaults were also added so that a predictable development database name could be used during local setup while still allowing institution-specific values in another environment.",
                    "PyMySQL is installed as the MySQLdb-compatible driver through a small initialization module, which allows Django’s MySQL backend to operate without altering higher-level application logic. Because contemporary MySQL servers frequently use sha256-based authentication modes, the cryptography package was added as a dependency to satisfy that requirement. This detail is operationally important and was observed during migration validation.",
                    "Migration refinement also revealed a historical inconsistency in the accounts migrations, where the department field already existed in the initial migration and a later migration attempted to add it again. Resolving that inconsistency by converting the redundant migration into a no-op helped the project align its migration history with the actual schema. This is precisely the kind of practical issue that surfaces when a project moves from a lightweight local database into a more disciplined relational environment.",
                ],
            },
        ],
    },
    {
        "number": 7,
        "title": "TIMETABLE GENERATION LOGIC",
        "start_page": 31,
        "intro": "This chapter explains the conflict-aware timetable engine that transforms subject and faculty records into weekly schedules and workload summaries.",
        "sections": [
            {
                "number": "7.1",
                "title": "Scheduling Inputs and Data Preparation",
                "paragraphs": [
                    "The timetable generator operates on subjects filtered by academic year and, when relevant, by department. This ensures that a schedule is produced from an academically coherent subset rather than from all subjects in the database at once. The selected records are grouped by semester so that multiple semester timetables can coexist in the same weekly output while still being scheduled independently at the slot level.",
                    "The project defines six teaching slots per day and a six-day academic week, with lunch treated as a fixed non-teaching interval. These values are represented explicitly in the generator module and are also mirrored in the dashboard and template layer so that the conceptual schedule and the visible schedule remain synchronized. Such explicitness is useful in academic software because hidden assumptions about time structure can easily lead to confusing outputs.",
                    "Before placement begins, the generator separates laboratory subjects from theory subjects using a simple but effective heuristic based on subject code or subject name. This preparation step is important because laboratory sessions require consecutive slot allocation, while theory subjects need distributed weekly placement. The quality of the final timetable depends heavily on handling these categories differently from the outset.",
                ],
            },
            {
                "number": "7.2",
                "title": "Conflict-Aware Placement Rules",
                "paragraphs": [
                    "Laboratory scheduling is performed first because it is more restrictive. Each lab needs a three-slot block either in the morning or in the afternoon. The algorithm checks that the target semester’s slots are free and also verifies that the same day-start lab block is not already occupied by another semester. This cross-semester protection reduces one of the most visible timetable conflicts in practice, namely simultaneous lab-block overlap across semester schedules.",
                    "Theory placement follows after laboratories are fixed. The algorithm shuffles the required subject-hour instances and attempts to place each one in a free slot while respecting two local rules: the same subject should not appear in consecutive slots for a semester, and the same subject should not repeat more than once on the same day for that semester. These constraints help spread learning activity more naturally across the week.",
                    "Although the algorithm uses randomness during shuffle order, it remains rule-aware rather than arbitrary. The practical effect is that the generator avoids many common manual mistakes without requiring an overly complex optimization engine. For a department-level academic portal, this balance between simplicity and usefulness is appropriate and easy to maintain.",
                ],
                "bullets": [
                    "Labs are allocated before theory subjects.",
                    "Lab blocks can start only in the morning or afternoon block boundaries.",
                    "A theory subject is not repeated consecutively for the same semester.",
                    "A theory subject is not repeated twice on the same day for the same semester.",
                ],
            },
            {
                "number": "7.3",
                "title": "Faculty Workload Calculation and Output Rendering",
                "paragraphs": [
                    "Once the timetable grid is populated, the generator computes faculty workload by traversing all scheduled cells. For each faculty code it records the faculty name, the number of theory hours, the number of practical hours and the set of subjects taught. The resulting structure supports both analytical reporting and interface-level personalization, such as highlighting the current faculty member’s row in the workload table.",
                    "In the rendered timetable view, semesters are displayed as rows within each day, lunch is visually separated and faculty-specific classes are highlighted in yellow when the logged-in user is a faculty member. This creates an immediate connection between backend scheduling logic and frontend usability. Faculty members do not need to scan the entire grid blindly; the interface guides them toward their own teaching commitments.",
                    "The workload summary also adds academic value beyond the timetable itself. Department assistants can use it to review balance across faculty members, and faculty users can verify whether the generated schedule reflects their expected teaching distribution. Thus, the algorithm produces both a weekly schedule and a management-friendly summary of teaching responsibility.",
                ],
            },
        ],
    },
    {
        "number": 8,
        "title": "MODULE DESCRIPTION AND USER INTERFACE",
        "start_page": 34,
        "intro": "This chapter documents the operational modules and the major user-facing screens implemented in the portal.",
        "sections": [
            {
                "number": "8.1",
                "title": "Authentication and Onboarding Module",
                "paragraphs": [
                    "The authentication module includes first-run setup, login, self-registration, sign-out, password reset generation and password reset completion. This combination allows the project to support both initial institutional bootstrap and ongoing user onboarding without manual database edits. A particularly useful academic feature is the requirement that most users must be pre-registered before they can self-register, which keeps the system closed to unauthorized public sign-ups.",
                    "Forms in this module are not passive data collectors. They enforce password confirmation, normalize email addresses, validate pre-registration status and prevent duplicate assistant assignment. The first-run form is deliberately separate because the very first main assistant must be created before any normal registration workflow exists. This design avoids the circular dependency that would occur if every account always required prior pre-registration.",
                    "From a user-interface perspective, the accounts screens communicate process clearly. The login page hosts both sign-in and registration flows, the forgot-password page shows a reset link in development mode, and the sign-out page adds a human confirmation step. These details strengthen the practical usability of the system for non-technical academic users.",
                ],
            },
            {
                "number": "8.2",
                "title": "Dashboard and Monitoring Module",
                "paragraphs": [
                    "The dashboard is the role-sensitive landing surface of the application. It greets the user by name, displays a role indicator and offers a summary of the actions most relevant to that role. Department assistants see counts of subjects, faculty, users and academic years for their department. The main assistant sees monitoring-oriented department statistics and read-only visibility across the academic structure.",
                    "This design is notable because it turns the dashboard into an interpretation layer rather than a decorative page. A user can understand immediately whether they are in a department-editing role, a university-monitoring role, a faculty-consumption role or a student-consumption role. The system also uses contextual banners to explain these distinctions, which reduces errors and improves confidence in the workflow.",
                    "The dashboard further includes a visual explanation of the daily schedule structure, including the lunch interval and the fact that lab subjects are allocated as three-hour blocks. This small but meaningful feature helps users interpret the generated timetable correctly before they open the full timetable page.",
                ],
            },
            {
                "number": "8.3",
                "title": "Academic Management Module",
                "paragraphs": [
                    "The academic management module consists of faculty management, subject management and user pre-registration. Faculty management allows department assistants to create and maintain instructor codes and names, with deletion blocked when a faculty record is still referenced by subjects. Subject management allows assistants to add academic years, map subjects to semesters and faculty, and set hours per week. These screens ensure that timetable inputs remain curated rather than improvised.",
                    "User management is especially important because it reflects the governance philosophy of the application. The main assistant can pre-register department assistants only, while department assistants can pre-register students and faculty within their own department. This rule is clearly expressed in the interface and enforced by the view logic, producing a healthy balance between institutional oversight and local ownership.",
                    "Taken together, these screens show that the portal is more than a viewer. It is a controlled academic administration environment in which core teaching data is prepared carefully before automated schedule generation begins. The project therefore models the preparatory work that real timetable generation depends on.",
                ],
            },
            {
                "number": "8.4",
                "title": "Resource Management and Timetable Presentation Module",
                "paragraphs": [
                    "The resource module allows department assistants to upload or replace semester-wise syllabus files and upload old question papers with subject and exam metadata. Students and faculty then browse these resources using department, year and semester filters. By integrating this feature directly into the same authenticated portal, the system reduces dependence on external file-sharing channels and provides a more complete academic service experience.",
                    "The timetable presentation module combines interactive filtering and print readiness. The main assistant can switch departments for monitoring purposes, while other users are locked to their own department. Academic years can be selected dynamically, and faculty users receive personal visual highlighting. Under the timetable grid, a workload table summarizes subject responsibility and total hours, which adds an analytical layer beyond mere viewing.",
                    "These two modules complement each other strongly. The timetable answers when classes happen, while the resource repository answers what documents students need for those classes. Their coexistence inside one portal makes the project more coherent and more valuable to academic stakeholders.",
                ],
            },
        ],
    },
    {
        "number": 9,
        "title": "TESTING, VALIDATION AND EXPERIMENTAL RESULTS",
        "start_page": 38,
        "intro": "This chapter discusses how the implemented system was checked, what issues were observed and what outcomes indicate successful operation.",
        "sections": [
            {
                "number": "9.1",
                "title": "Testing Strategy",
                "paragraphs": [
                    "Testing in this project combines automated checks, migration verification, manual role-flow inspection and output-oriented timetable review. The automated test file currently focuses on the account-switching workflow, proving that a logged-in assistant can sign out and then log in again as another role. Even though this test suite is small, it demonstrates how important session-handling behavior can be captured formally.",
                    "Migration and server startup testing were especially relevant during the MySQL transition. Practical issues such as missing cryptography support, nonexistent database creation and duplicate migration behavior were discovered during real execution. Solving these issues improved the reliability of the project and also provided useful documentation evidence that the software was exercised under realistic deployment conditions rather than being described only at a theoretical level.",
                    "Manual role-flow testing remains essential because many system guarantees depend on authorization behavior and filtered querysets. For example, department assistants must be unable to access other departments’ data, and the main assistant must remain monitor-only for data-entry pages. These properties are best confirmed by interacting with the system from multiple role perspectives.",
                ],
            },
            {
                "number": "9.2",
                "title": "Representative Test Cases",
                "paragraphs": [
                    "A useful way to evaluate the project is through representative academic scenarios. The first scenario verifies bootstrap behavior by confirming that, when no user exists, the platform redirects to first-run setup and allows the main assistant account to be created. The second scenario verifies pre-registration by ensuring that an unregistered email cannot self-register but an authorized email can do so successfully.",
                    "A third scenario concerns timetable data preparation. Here the department assistant adds faculty and subjects for an academic year and semester, associates each subject with a faculty record and then opens the timetable page to generate output. The expected result is a semester-wise timetable accompanied by faculty workload data. A fourth scenario validates resource handling by uploading a syllabus and a question paper and then confirming that the student-facing resources page can retrieve the correct records using year and semester filters.",
                    "These scenarios matter because they span the system from account creation to academic output. A project that passes all of them is demonstrating not a narrow feature success but a coherent workflow success across the institution’s daily academic needs.",
                ],
                "table": {
                    "headers": ["Test Case", "Expected Result", "Observed Outcome"],
                    "rows": [
                        ["First-run setup", "First main assistant account created successfully", "Supported by dedicated first-run form and redirect logic"],
                        ["Pre-registered self-registration", "Authorized user creates account; unauthorized email is rejected", "Handled through form validation and slot checks"],
                        ["Timetable generation", "Department schedule and workload table appear for selected year", "Generated from filtered subjects and faculty relations"],
                        ["Resource upload and retrieval", "Uploaded files become available with metadata filters", "Implemented through syllabus and question-paper models"],
                    ],
                },
            },
            {
                "number": "9.3",
                "title": "Observed Issues and Corrective Actions",
                "paragraphs": [
                    "The MySQL migration process surfaced several instructive issues. Initially the server connection required cryptography support because of the authentication method used by the local MySQL installation. This was resolved by adding the dependency explicitly. Next, the target database needed to be created before migrations could run. Finally, the historical duplication of the department field in the accounts migration chain caused a duplicate-column error and had to be neutralized by converting the redundant migration into an empty operation.",
                    "These issues are valuable from a project-learning perspective. They show that application development is not finished when the code appears logically correct inside a repository. Real deployment work often reveals missing dependencies, schema-history inconsistencies and environmental assumptions that must be addressed before software becomes truly usable.",
                    "The corrective actions taken were disciplined and minimal. Dependencies were updated, settings were aligned with MySQL, and the migration chain was repaired without rewriting unrelated application logic. This demonstrates good maintenance practice and provides a concrete example of how a student project evolves toward a cleaner deployment state.",
                ],
            },
            {
                "number": "9.4",
                "title": "Experimental Results and Interpretation",
                "paragraphs": [
                    "The experimental result of the project is a functioning academic portal that can authenticate multiple roles, constrain them according to institutional responsibility, generate a rule-aware timetable, summarize faculty workload and serve academic documents through department-aware filtering. These results matter because they address the real administrative pain points identified earlier in the report rather than solving a purely abstract scheduling puzzle.",
                    "From an implementation perspective, the results demonstrate stable interaction between the accounts and core modules. User registration feeds into access control, faculty and subject records feed into timetable generation, and the timetable output feeds into faculty-specific highlighting and workload summaries. The resource module broadens the system’s utility further by serving syllabi and previous papers from the same database-backed environment.",
                    "The project therefore succeeds not only by running but by producing integrated academic behavior. The final result is a more organized, more governable and more extensible platform than the manual or spreadsheet-centered processes it is designed to improve.",
                ],
            },
        ],
    },
    {
        "number": 10,
        "title": "DEPLOYMENT, MAINTENANCE AND CONCLUSION",
        "start_page": 42,
        "intro": "This chapter summarizes operational deployment considerations, ongoing maintenance concerns, conclusions and future enhancement opportunities.",
        "sections": [
            {
                "number": "10.1",
                "title": "Deployment and Operational Readiness",
                "paragraphs": [
                    "The documented project is ready for controlled local deployment with MySQL and a Django development server, and it can be prepared for broader hosting with relatively small additional effort. The essential steps are dependency installation, MySQL database creation, environment-variable configuration, migration execution and server startup. Because the application already separates static and media behavior and stores academic files with metadata, it is structurally well positioned for transition to a more persistent hosting environment.",
                    "Operational readiness also depends on account policy and backup practice. The first-run setup mechanism ensures a clean administrative bootstrap, while environment-based database credentials reduce the need to hard-code sensitive values. Since the application stores user identities and uploaded academic files, regular database and media backups should be treated as part of the deployment plan from the beginning.",
                    "The MySQL configuration makes the project more realistic for deployment than a SQLite-only prototype. Multi-user administrative workflows, central backup routines and server-oriented hosting all align more naturally with a dedicated MySQL database. This strengthens the practical value of the project for academic demonstration and future institutional use.",
                ],
            },
            {
                "number": "10.2",
                "title": "Maintenance and Future Scope",
                "paragraphs": [
                    "Ongoing maintenance for the project will likely focus on three areas: data quality, user support and feature growth. Data quality maintenance includes reviewing subject assignments, retiring obsolete academic years and ensuring uploaded resources remain current. User support maintenance includes password recovery, account-role corrections and assistant turnover management. Feature growth may include classroom allocation, email notifications, richer analytics, attendance integration and export workflows.",
                    "The codebase is prepared reasonably well for such maintenance because the accounts and core responsibilities are separated and because the most important institutional rules are expressed in models, forms and decorators rather than being hidden inside templates. This means future developers can extend the system while still understanding where role policy, domain logic and presentation logic belong.",
                    "A particularly promising future direction is stronger optimization in the timetable engine. The present rule-aware random placement approach is practical for department-level scheduling, but future work could incorporate room allocation, faculty preferences, immutable blocked periods or deterministic optimization strategies. These improvements can build on the existing data model rather than replacing it entirely.",
                ],
            },
            {
                "number": "10.3",
                "title": "Overall Conclusion",
                "paragraphs": [
                    "The Conflict Aware Timetable Management System demonstrates how a university-oriented web application can combine academic administration, access control, document sharing and timetable generation inside a single coherent platform. Its value lies not merely in producing a weekly schedule, but in organizing the full workflow that makes such a schedule trustworthy and usable.",
                    "The project’s transition to MySQL is a meaningful part of that conclusion. It shows that the application is capable of moving beyond a local prototype database into a more credible multi-user environment. Alongside the resolved migration issues and the implemented access-control structure, this indicates a healthy level of engineering maturity for an academic project.",
                    "In summary, the system provides a strong functional foundation, a clean modular structure and a realistic administrative model. It can already serve as a useful departmental portal, and it also provides substantial room for future enhancement in scheduling intelligence, reporting depth and production deployment sophistication.",
                ],
            },
        ],
    },
]


REFERENCES = [
    "Django Software Foundation. Django Documentation for Version 4.2.",
    "Oracle Corporation. MySQL Reference Manual.",
    "Python Software Foundation. Python Language Reference and Standard Library Documentation.",
    "Bootstrap Team. Bootstrap User Interface Documentation.",
    "Silberschatz, Korth and Sudarshan. Database System Concepts.",
    "Ian Sommerville. Software Engineering.",
    "Roger S. Pressman and Bruce R. Maxim. Software Engineering: A Practitioner’s Approach.",
    "Ralph Stair and George Reynolds. Principles of Information Systems.",
    "Academic scheduling and timetabling research literature on conflict-aware allocation and constraint satisfaction.",
    "Project source code and runtime configuration in the Conflict Aware Timetable Management System repository.",
]


APPENDIX_FILES = [
    ("Appendix A.1 - User and Pre-Registration Models", ROOT / "accounts" / "models.py", 1, 160),
    ("Appendix A.2 - Core Academic Models", ROOT / "core" / "models.py", 1, 130),
    ("Appendix A.3 - Timetable Generation Module", ROOT / "core" / "timetable.py", 1, 150),
]


TABLE_COUNTER = 0


def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def configure_page(section):
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)
    section.page_width = Mm(210)
    section.page_height = Mm(297)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_font(run, size=12)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_text = OxmlElement("w:t")
    fld_text.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(fld_text)
    run._r.append(fld_end)


def ensure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.first_line_indent = Inches(0.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.left_indent = Cm(0)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.space_before = Pt(12 if style_name == "Heading 1" else 6)

    doc.styles["Heading 1"].font.size = Pt(14)
    doc.styles["Heading 2"].font.size = Pt(12)
    doc.styles["Heading 3"].font.size = Pt(12)

    if "CodeBlock" not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Courier New"
        code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        code_style.font.size = Pt(9.5)
        code_pf = code_style.paragraph_format
        code_pf.line_spacing = 1.0
        code_pf.space_before = Pt(0)
        code_pf.space_after = Pt(0)
        code_pf.left_indent = Cm(0)
        code_pf.right_indent = Cm(0)
        code_pf.first_line_indent = Cm(0)


def add_body_paragraph(doc: Document, text: str, first_line=True):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5) if first_line else Cm(0)
    run = p.add_run(text)
    set_font(run, size=12)
    return p


def add_heading1(doc: Document, text: str):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_font(run, size=14, bold=True, color=(0, 0, 0))
    return p


def add_heading2(doc: Document, text: str):
    p = doc.add_paragraph(style="Heading 2")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color=(0, 0, 0))
    return p


def add_bullet_list(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run("• ").bold = True
        run = p.add_run(item)
        set_font(run, size=12)


def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_font(run, size=11, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str | None = None):
    global TABLE_COUNTER
    if caption:
        TABLE_COUNTER += 1
        cap = doc.add_paragraph(style="Normal")
        cap.paragraph_format.first_line_indent = Cm(0)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"Table {TABLE_COUNTER}: {caption}")
        set_font(run, size=12, bold=True)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, value in enumerate(headers):
        set_cell(hdr[idx], value, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(value) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cells[idx], value, bold=False, align=align)

    doc.add_paragraph("")


def add_text_matrix(doc: Document, headers: list[str], rows: list[list[str]], caption: str | None = None):
    global TABLE_COUNTER
    if caption:
        TABLE_COUNTER += 1
        cap = doc.add_paragraph(style="Normal")
        cap.paragraph_format.first_line_indent = Cm(0)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"Table {TABLE_COUNTER}: {caption}")
        set_font(run, size=12, bold=True)

    for row in rows:
        parts = [f"{headers[idx]} - {row[idx]}" for idx in range(1, len(row))]
        add_bullet_list(doc, [f"{row[0]}: {'; '.join(parts)}"])


def center_page(doc: Document, lines: list[tuple[str, int, bool]], add_break=True):
    for text, size, bold in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        set_font(run, size=size, bold=bold)
    if add_break:
        doc.add_page_break()


def add_cover_page(doc: Document):
    lines = [
        ("A MAJOR PROJECT WORK SUBMITTED FOR THE PARTIAL FULFILLMENT", 12, False),
        ("FOR THE AWARD OF THE DEGREE OF", 12, False),
        ("", 12, False),
        ("MASTER OF COMPUTER APPLICATIONS", 14, True),
        ("", 12, False),
        (PROJECT_OVERVIEW["title"], 18, True),
        ("", 12, False),
        ("Submitted By", 12, True),
        ("SHANTHIPRIYA JANA", 13, True),
        ("", 12, False),
        ("Under the esteemed guidance of", 12, False),
        ("PROJECT GUIDE", 12, True),
        ("", 12, False),
        ("Department of Computer Science", 12, True),
        ("VIKRAMA SIMHAPURI UNIVERSITY", 13, True),
        ("NELLORE - 524 324", 12, False),
        ("www.vsu.ac.in", 12, False),
        ("", 12, False),
        ("April - 2026", 12, True),
    ]
    center_page(doc, lines)


def add_certificate_page(doc: Document):
    center_page(
        doc,
        [
            ("VIKRAMA SIMHAPURI UNIVERSITY", 13, True),
            ("DEPARTMENT OF COMPUTER SCIENCE", 12, True),
            ("NELLORE", 12, True),
            ("", 12, False),
            ("CERTIFICATE", 14, True),
        ],
        add_break=False,
    )
    add_body_paragraph(
        doc,
        "This is to certify that the major project work entitled \"Conflict Aware Timetable Management System\" submitted to Vikrama Simhapuri University is a record of bonafide work carried out for the study, analysis, design and implementation of a role-based academic scheduling and resource management portal using Django and MySQL. The work embodied in this report has been prepared in a formal academic format and represents a project-specific documentation of the implemented system.",
        first_line=False,
    )
    add_body_paragraph(
        doc,
        "The report describes the implemented Django project available in the current workspace and explains the manner in which the system supports first-run setup, controlled self-registration, faculty and subject administration, resource uploads, timetable generation and workload reporting. It also records the migration of the application from SQLite-style development assumptions to a MySQL-backed execution model appropriate for multi-user academic environments.",
    )
    add_body_paragraph(
        doc,
        "The documentation has been prepared in a structured professional format so that the project can be evaluated not only as software, but also as a complete academic deliverable. All major modules, workflows, constraints and deployment considerations have been explained with due attention to clarity, organization and technical accuracy.",
    )
    add_body_paragraph(doc, "(PROJECT GUIDE)                                            (HEAD OF THE DEPARTMENT)", first_line=False)
    doc.add_page_break()


def add_declaration_page(doc: Document):
    add_heading1(doc, "DECLARATION")
    add_body_paragraph(
        doc,
        "This documentation is prepared as an original academic report for the implemented Conflict Aware Timetable Management System available in the working project directory. The report consolidates the actual structure, models, views, templates, database configuration and scheduling logic present in the project and does not merely describe a hypothetical application. All technical descriptions are based on the current state of the codebase and the MySQL-oriented execution flow used for this version.",
        first_line=False,
    )
    add_body_paragraph(
        doc,
        "Where interpretive analysis has been added, it has been done to explain the purpose and significance of the implemented modules in a university administration context. The intention of the report is to provide a coherent, long-form academic record of the system, its rationale, its design and its practical deployment considerations.",
    )
    add_body_paragraph(
        doc,
        "This declaration page is included to reflect the formal style usually expected in a major project document and to establish that the report is prepared specifically for the present software implementation rather than adapted from an unrelated template.",
    )
    doc.add_page_break()


def add_acknowledgement_page(doc: Document):
    add_heading1(doc, "ACKNOWLEDGMENT")
    add_body_paragraph(
        doc,
        "The preparation of this documentation benefited from the structure already present in the project source code, including its modular Django layout, role-aware views, data models and templates. A well-organized codebase makes it possible to write documentation that is not generic but grounded in real implementation choices. For that reason, the source files themselves deserve acknowledgement as the primary technical reference for this report.",
        first_line=False,
    )
    add_body_paragraph(
        doc,
        "The project also reflects the practical needs of academic administration, especially the desire to reduce timetable conflicts, provide clear departmental ownership of data and create a more convenient channel for students to access syllabi and previous question papers. These institutional needs shaped the structure of the application and gave the documentation a clear functional direction.",
    )
    add_body_paragraph(
        doc,
        "Finally, acknowledgement is due to the development and testing process itself. The move to MySQL, the handling of migration issues and the validation of user roles helped transform the project from a local prototype into a more deployment-oriented system. Those refinements significantly strengthened the quality and relevance of the final document.",
    )
    doc.add_page_break()


def add_abstract_page(doc: Document):
    center_page(
        doc,
        [
            (PROJECT_OVERVIEW["title"], 14, True),
            ("", 12, False),
            ("ABSTRACT", 14, True),
        ],
        add_break=False,
    )
    abstract = (
        "The Conflict Aware Timetable Management System is a web-based academic administration platform developed with Django and adapted for MySQL-backed operation. The system is designed to support university or department workflows by combining role-based user control, subject and faculty management, automated timetable generation, workload reporting and academic resource sharing. A main assistant can initialize and monitor the platform, department assistants can manage departmental data, faculty members can review personalized schedules and students can access timetable and learning resources. The timetable engine applies practical scheduling rules such as semester grouping, laboratory block allocation, non-consecutive repetition control and daily distribution of theory classes. In addition to timetable functions, the application includes self-registration through pre-authorized email records, password reset tokens, syllabus uploads and old question paper management. This documentation presents the complete project rationale, methodology, architecture, data model, implementation details, MySQL configuration, testing observations and future enhancement scope in a formal academic format."
    )
    add_body_paragraph(doc, abstract, first_line=False)
    doc.add_page_break()


def add_toc_page(doc: Document):
    add_heading1(doc, "TABLE OF CONTENTS")
    toc_lines = [
        "ABSTRACT - Page 05",
        "1. INTRODUCTION - Page 09",
        "2. LITERATURE REVIEW - Page 13",
        "3. METHODOLOGY - Page 16",
        "4. SYSTEM REQUIREMENTS AND SPECIFICATION - Page 20",
        "5. SYSTEM ANALYSIS AND ARCHITECTURE - Page 23",
        "6. DATABASE DESIGN AND MYSQL IMPLEMENTATION - Page 27",
        "7. TIMETABLE GENERATION LOGIC - Page 31",
        "8. MODULE DESCRIPTION AND USER INTERFACE - Page 34",
        "9. TESTING, VALIDATION AND EXPERIMENTAL RESULTS - Page 38",
        "10. DEPLOYMENT, MAINTENANCE AND CONCLUSION - Page 42",
        "11. REFERENCES - Page 45",
        "12. APPENDIX A: SELECTED SOURCE CODE - Page 46",
    ]
    add_bullet_list(doc, toc_lines)
    doc.add_page_break()


def add_lists_page(doc: Document):
    add_heading1(doc, "LIST OF TABLES AND ABBREVIATIONS")
    add_body_paragraph(
        doc,
        "This project report includes descriptive tables for objectives, architecture, role permissions, database entities, software requirements and representative test cases. These tables are used to improve scanability and to present the system structure in an academically familiar way.",
        first_line=False,
    )
    add_heading2(doc, "Common Abbreviations")
    add_bullet_list(
        doc,
        [
            "DBMS - Database Management System",
            "ORM - Object Relational Mapper",
            "UI - User Interface",
            "URL - Uniform Resource Locator",
            "OTP-style token flow - Single-use password reset token approach",
            "CSE - Computer Science and Engineering",
            "API - Application Programming Interface",
            "RDBMS - Relational Database Management System",
        ],
    )
    doc.add_page_break()


def add_inventory_page(doc: Document):
    model_count = 7
    template_count = len(list((ROOT / "templates").rglob("*.html")))
    py_files = len(list(ROOT.rglob("*.py")))
    add_heading1(doc, "PROJECT INVENTORY")
    add_body_paragraph(
        doc,
        "Before the chapter-wise narrative begins, this page summarizes a few measurable characteristics of the implementation. These values help situate the scope of the project and make it easier to interpret later discussions on architecture, module boundaries and source-code appendices.",
        first_line=False,
    )
    add_text_matrix(
        doc,
        ["Metric", "Observed Value", "Interpretation"],
        [
            ["Python source files", str(py_files), "Indicates the breadth of backend logic and configuration assets."],
            ["HTML templates", str(template_count), "Shows the number of user-facing views prepared in the template layer."],
            ["Core persisted entities", str(model_count), "Represents the principal academic and user-management records."],
            ["Major Django apps", "2", "Accounts and Core separate identity from academic operations."],
        ],
        caption="High-level implementation inventory",
    )
    doc.add_page_break()


def add_chapter_sections(doc: Document):
    for chapter in CHAPTERS:
        first_in_chapter = True
        for idx, section in enumerate(chapter["sections"]):
            if first_in_chapter:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(40)
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(chapter["title"])
                set_font(run, size=14, bold=True, color=(0, 0, 0))
                add_body_paragraph(doc, chapter["intro"], first_line=False)
                first_in_chapter = False
            add_heading2(doc, f'{section["number"]} {section["title"]}')
            for para_idx, para in enumerate(section["paragraphs"]):
                add_body_paragraph(doc, para, first_line=(para_idx != 0))
            if section.get("bullets"):
                add_bullet_list(doc, section["bullets"])
            if section.get("table"):
                add_text_matrix(
                    doc,
                    section["table"]["headers"],
                    section["table"]["rows"],
                    caption=section["title"],
                )
            if not (chapter == CHAPTERS[-1] and idx == len(chapter["sections"]) - 1):
                doc.add_page_break()


def add_references_page(doc: Document):
    doc.add_page_break()
    add_heading1(doc, "REFERENCES")
    for i, item in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.hanging_indent = Inches(0.3)
        run = p.add_run(f"{i}. {item}")
        set_font(run, size=12)


def split_wrapped_code(text: str, width: int = 84) -> list[str]:
    wrapped_lines = []
    for line in text.splitlines():
        if not line.strip():
            wrapped_lines.append("")
            continue
        indent = len(line) - len(line.lstrip(" "))
        initial = " " * indent
        wrapped = textwrap.wrap(
            line,
            width=width,
            subsequent_indent=initial + "    ",
            replace_whitespace=False,
            drop_whitespace=False,
        )
        wrapped_lines.extend(wrapped or [line])
    return wrapped_lines


def add_code_paragraph(doc: Document, line: str):
    p = doc.add_paragraph(style="CodeBlock")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(line)
    set_font(run, name="Courier New", size=9.5)


def add_appendices(doc: Document):
    doc.add_page_break()
    add_heading1(doc, "APPENDIX A: SELECTED SOURCE CODE")
    add_body_paragraph(
        doc,
        "The appendix includes only a small set of representative source-code excerpts so that the documentation stays focused on analysis, design and implementation discussion. The selected snippets show the most important model and timetable logic without turning the report into a large code archive.",
        first_line=False,
    )
    for title, path, start_line, end_line in APPENDIX_FILES:
        doc.add_page_break()
        add_heading2(doc, title)
        add_body_paragraph(doc, f"Source file: {path.relative_to(ROOT)}", first_line=False)
        lines = path.read_text(encoding="utf-8").splitlines()
        selected = "\n".join(lines[start_line - 1:end_line])
        text = selected
        wrapped_lines = split_wrapped_code(text)
        chunk_size = 110
        for start in range(0, len(wrapped_lines), chunk_size):
            chunk = wrapped_lines[start:start + chunk_size]
            if start > 0:
                doc.add_page_break()
                add_heading2(doc, f"{title} (continued)")
                add_body_paragraph(doc, f"Source file: {path.relative_to(ROOT)}", first_line=False)
            for offset, line in enumerate(chunk, start=start):
                logical_line = start_line + offset
                add_code_paragraph(doc, f"{logical_line:04d}: {line}")


def build_document():
    doc = Document()
    configure_page(doc.sections[0])
    ensure_styles(doc)
    for section in doc.sections:
        configure_page(section)
        footer_p = section.footer.paragraphs[0]
        footer_p.text = ""
        add_page_number(footer_p)

    add_cover_page(doc)
    add_certificate_page(doc)
    add_declaration_page(doc)
    add_acknowledgement_page(doc)
    add_abstract_page(doc)
    add_toc_page(doc)
    add_lists_page(doc)
    add_inventory_page(doc)
    add_chapter_sections(doc)
    add_references_page(doc)
    add_appendices(doc)
    doc.save(OUTPUT)
    print(f"Saved documentation to {OUTPUT}")


if __name__ == "__main__":
    build_document()
